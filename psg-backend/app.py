import sqlite3
import os
import io
import json
from flask import Flask, jsonify, request, send_file
from flask_cors import CORS
from docx import Document
from docx.shared import RGBColor, Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from datetime import datetime
import tempfile
import smtplib
from email.mime.text import MIMEText
import requests
from pathlib import Path

# Configuração da aplicação Flask
app = Flask(__name__)
CORS(app)

# --- Funções de Banco de Dados ---
def get_db_connection():
    """Cria e retorna a conexão com o banco de dados."""
    conn = sqlite3.connect(DATABASE)
    conn.row_factory = sqlite3.Row
    return conn

def obter_proximo_numero(pasta_nome):
    """Retorna o próximo número disponível para o documento na pasta."""
    full_path = os.path.join(BASE_PATH, pasta_nome)

    if not os.path.exists(full_path):
        return "01"

    conn = get_db_connection()
    sigla_pasta = conn.execute("SELECT sigla FROM pastas WHERE nome = ?", (pasta_nome,)).fetchone()['sigla']
    conn.close()

    arquivos_existentes = [f for f in os.listdir(full_path) if f.startswith(f"PSG-{sigla_pasta}-")]
    numeros = []
    
    for arquivo in arquivos_existentes:
        try:
            partes = arquivo.split('-')
            if len(partes) >= 3:
                numero_str = partes[-1].split('.')[0]
                if numero_str.isdigit():
                    numeros.append(int(numero_str))
        except:
            continue
    
    proximo_numero = max(numeros) + 1 if numeros else 1
    return f"{proximo_numero:02d}"

# --- Funções de Geração de Documento ---
def criar_cabecalho(doc, titulo_pasta, sigla_pasta, tema_sigla, numero_documento):
    """Cria e configura o cabeçalho do documento."""
    section = doc.sections[0]
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)

    header = section.header
    for paragraph in header.paragraphs:
        paragraph.clear()
    
    tabela = header.add_table(rows=1, cols=3, width=Inches(7.5))
    colunas = tabela.columns
    colunas[0].width = Inches(0.2)
    colunas[1].width = Inches(0.2)
    colunas[2].width = Inches(0.2)

    tbl = tabela._tbl
    tblPr = tbl.tblPr
    tblCellMar = OxmlElement('w:tblCellMar')
    for tag in ['top', 'left', 'bottom', 'right']:
        elem = OxmlElement(f'w:{tag}')
        elem.set(qn('w:w'), '0')
        elem.set(qn('w:type'), 'dxa')
        tblCellMar.append(elem)
    tblPr.append(tblCellMar)

    linha = tabela.rows[0].cells
    
    try:
        para_img_esq = linha[0].paragraphs[0]
        run_img_esq = para_img_esq.add_run()
        run_img_esq.add_picture(IMAGEM_ESQUERDA, width=Inches(0.8))
        para_img_esq.alignment = WD_ALIGN_PARAGRAPH.LEFT
        para_img_esq.paragraph_format.space_before = Pt(0)
        para_img_esq.paragraph_format.space_after = Pt(0)
    except Exception as e:
        print(f"Erro ao adicionar imagem da esquerda: {e}")

    conteudo_central = linha[1].paragraphs[0]
    run_titulo = conteudo_central.add_run(f"PSG - {titulo_pasta} - {tema_sigla}\n")
    run_titulo.bold = True
    run_titulo.font.size = Pt(12)
    
    run_linha = conteudo_central.add_run("_" * 99 + "\n")
    run_linha.font.size = Pt(10)
    run_linha.font.bold = True
    
    codigo_psg = f"PSG.{sigla_pasta}.{tema_sigla}.{numero_documento}"
    data_aprovacao = datetime.now().strftime("%d/%m/%Y")
    texto_cabecalho = f"{codigo_psg}\t| Rev.\t{numero_documento}\t| Aprovação:\t{data_aprovacao}|"
    run_info = conteudo_central.add_run(texto_cabecalho)
    run_info.font.size = Pt(10)
    run_info.bold = True
    conteudo_central.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    try:
        para_img_dir = linha[2].paragraphs[0]
        run_img_dir = para_img_dir.add_run()
        run_img_dir.add_picture(IMAGEM_DIREITA, width=Inches(0.6))
        para_img_dir.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        para_img_dir.paragraph_format.space_before = Pt(0)
        para_img_dir.paragraph_format.space_after = Pt(0)
    except Exception as e:
        print(f"Erro ao adicionar imagem da direita: {e}")
        
    return doc

def adicionar_secao(doc, titulo_secao, conteudo):
    """Adiciona uma seção com título e conteúdo ao documento."""
    doc.add_heading(titulo_secao, level=2)
    doc.add_paragraph(conteudo if conteudo.strip() else "Não informado.")

def send_notification_email(destinatario, assunto, corpo):
    """Envia um email de notificação."""
    try:
        msg = MIMEText(corpo)
        msg['Subject'] = assunto
        msg['From'] = EMAIL_USER
        msg['To'] = destinatario

        with smtplib.SMTP(EMAIL_HOST, EMAIL_PORT) as server:
            server.starttls()
            server.login(EMAIL_USER, EMAIL_PASS)
            server.send_message(msg)
        print("✅ E-mail de notificação enviado com sucesso!")
        return True
    except Exception as e:
        print(f"❌ Erro ao enviar e-mail: {e}")
        return False

def check_and_create_sharepoint_folder(folder_name):
    """
    Verifica se uma pasta existe no SharePoint e a cria se não existir.
    Retorna True se a pasta existir ou for criada, False caso contrário.
    """
    try:
        token_url = f'https://login.microsoftonline.com/{SHAREPOINT_TENANT_ID}/oauth2/v2.0/token'
        token_data = {
            'grant_type': 'client_credentials',
            'client_id': SHAREPOINT_CLIENT_ID,
            'client_secret': SHAREPOINT_CLIENT_SECRET,
            'scope': 'https://graph.microsoft.com/.default'
        }
        token_response = requests.post(token_url, data=token_data)
        access_token = token_response.json().get('access_token')

        headers = {'Authorization': f'Bearer {access_token}'}

        domain = SHAREPOINT_SITE_URL.split('/')[2]
        site_path = '/' + '/'.join(SHAREPOINT_SITE_URL.split('/')[3:])
        site_id_url = f'https://graph.microsoft.com/v1.0/sites/{domain}:{site_path}'
        site_id = requests.get(site_id_url, headers=headers).json().get('id')

        drive_url = f'https://graph.microsoft.com/v1.0/sites/{site_id}/drives'
        drive_response = requests.get(drive_url, headers=headers)
        drives = drive_response.json().get('value', [])
        drive_id = next((d['id'] for d in drives if d['name'] == 'Documentos'), None)
        
        if not drive_id:
            print("❌ Drive 'Documentos' não encontrado.")
            return False

        caminho_completo = f'{SHAREPOINT_BASE_PATH}/{folder_name}'
        check_url = f'https://graph.microsoft.com/v1.0/drives/{drive_id}/root:/{caminho_completo}'
        check_response = requests.get(check_url, headers=headers)

        if check_response.status_code == 200:
            print(f"✅ A pasta '{caminho_completo}' já existe no SharePoint.")
            return True
        elif check_response.status_code == 404:
            # A pasta não existe, vamos criá-la
            create_url = f'https://graph.microsoft.com/v1.0/drives/{drive_id}/root:/{caminho_completo}'
            create_payload = {'name': folder_name, 'folder': {}}
            create_response = requests.patch(create_url, headers={
                'Authorization': f'Bearer {access_token}',
                'Content-Type': 'application/json'
            }, data=json.dumps(create_payload))
            if create_response.status_code in [200, 201]:
                print(f"✅ Pasta '{caminho_completo}' criada no SharePoint!")
                return True
            else:
                print(f"❌ Erro ao criar pasta no SharePoint. Status: {create_response.status_code} \n\n {create_response.text}")
                return False
        else:
            print(f"❌ Erro ao verificar pasta no SharePoint. Status: {check_response.status_code} \n\n {check_response.text}")
            return False

    except Exception as e:
        print(f"❌ Erro na função check_and_create_sharepoint_folder: {e}")
        return False


def send_to_sharepoint(file_path, file_name, pasta_destino):
    """Envia o arquivo para uma pasta no SharePoint."""
    try:
        token_url = f'https://login.microsoftonline.com/{SHAREPOINT_TENANT_ID}/oauth2/v2.0/token'
        token_data = {
            'grant_type': 'client_credentials',
            'client_id': SHAREPOINT_CLIENT_ID,
            'client_secret': SHAREPOINT_CLIENT_SECRET,
            'scope': 'https://graph.microsoft.com/.default'
        }
        token_response = requests.post(token_url, data=token_data)
        access_token = token_response.json().get('access_token')

        headers = {'Authorization': f'Bearer {access_token}'}

        domain = SHAREPOINT_SITE_URL.split('/')[2]
        site_path = '/' + '/'.join(SHAREPOINT_SITE_URL.split('/')[3:])
        site_id_url = f'https://graph.microsoft.com/v1.0/sites/{domain}:{site_path}'
        site_id = requests.get(site_id_url, headers=headers).json().get('id')

        drive_url = f'https://graph.microsoft.com/v1.0/sites/{site_id}/drives'
        drive_response = requests.get(drive_url, headers=headers)
        drives = drive_response.json().get('value', [])
        drive_id = next((d['id'] for d in drives if d['name'] == 'Documentos'), None)
        
        if not drive_id:
            print("❌ Drive 'Documentos' não encontrado.")
            return False

        caminho_completo = f'{SHAREPOINT_BASE_PATH}/{pasta_destino}'
        
        upload_url = f'https://graph.microsoft.com/v1.0/drives/{drive_id}/root:/{caminho_completo}/{file_name}:/content'
        
        with open(file_path, 'rb') as file_content:
            upload_response = requests.put(upload_url, headers={
                'Authorization': f'Bearer {access_token}',
                'Content-Type': 'application/octet-stream'
            }, data=file_content)

        if upload_response.status_code in [200, 201]:
            print('✅ Arquivo enviado para o SharePoint com sucesso!')
            return True
        else:
            print(f"❌ Erro ao enviar arquivo para o SharePoint. Status: {upload_response.status_code} \n\n {upload_response.text}")
            return False

    except Exception as e:
        print(f"❌ Erro na função send_to_sharepoint: {e}")
        return False

# --- Rotas da API ---
@app.route('/pastas', methods=['GET'])
def get_pastas():
    """Retorna a lista de pastas e siglas do banco de dados."""
    conn = get_db_connection()
    pastas = conn.execute("SELECT id, nome, sigla FROM pastas").fetchall()
    conn.close()
    
    pastas_list = [{'id': p['id'], 'nome': p['nome'], 'sigla': p['sigla']} for p in pastas]
    return jsonify(pastas_list)

@app.route('/criar_pasta', methods=['POST'])
def criar_pasta():
    """
    Recebe nome e sigla para criar uma nova pasta.
    Verifica no SharePoint antes de salvar localmente.
    """
    data = request.json
    nome = data.get('nome')
    sigla = data.get('sigla')
    
    if not nome or not sigla:
        return jsonify({"error": "Nome e sigla são obrigatórios"}), 400

    # 1. Verifica e cria a pasta no SharePoint
    sharepoint_pasta_nome = nome
    if not check_and_create_sharepoint_folder(sharepoint_pasta_nome):
        # Se a pasta já existe no SharePoint, retorna um erro amigável
        return jsonify({"error": f"Pasta '{nome}' já existe no SharePoint."}), 409

    # 2. Se a pasta for criada com sucesso no SharePoint, salva localmente
    conn = get_db_connection()
    try:
        folder_path = os.path.join(BASE_PATH, nome)
        os.makedirs(folder_path, exist_ok=True)
        
        cursor = conn.cursor()
        cursor.execute("INSERT INTO pastas (nome, sigla) VALUES (?, ?)", (nome, sigla))
        conn.commit()
        
        novo_id = cursor.lastrowid
        
        return jsonify({"message": f"Pasta '{nome}' criada com sucesso!", "id": novo_id}), 201
    except sqlite3.IntegrityError:
        # Se a sigla já existe, remove o diretório criado
        os.rmdir(folder_path)
        return jsonify({"error": "Erro: Esta sigla já existe!"}), 409
    except Exception as e:
        return jsonify({"error": str(e)}), 500
    finally:
        conn.close()

@app.route('/gerar_documento', methods=['POST'])
def gerar_documento():
    """Recebe os dados do formulário, gera o documento e o envia para download."""
    try:
        data = request.form
        
        titulo = data.get('titulo', '')
        tema_sigla = data.get('tema', '')
        folder = json.loads(data.get('folder', '{}'))
        revisoes = json.loads(data.get('revisoes', '[]'))
        email = data.get('email', '')
        
        objetivo = data.get('objetivo', '')
        responsabilidades = data.get('responsabilidades', '')
        conceitosSiglas = data.get('conceitosSiglas', '')
        diretrizes = data.get('diretrizes', '')
        documentosComplementares = data.get('documentosComplementares', '')
        referencias = data.get('referencias', '')
        
        pasta_nome = folder.get('nome', '')
        sigla_pasta = folder.get('sigla', 'XXX')
        pasta_id = folder.get('id')
        
        if not pasta_nome or not pasta_id:
            return jsonify({"error": "Pasta não selecionada ou inválida."}), 400

        conn = get_db_connection()
        last_id = conn.execute("SELECT id FROM documentos ORDER BY id DESC LIMIT 1").fetchone()
        numero_documento = f"{(last_id['id'] if last_id else 0) + 1:02d}"
        
        doc = Document()
        doc = criar_cabecalho(doc, pasta_nome, sigla_pasta, tema_sigla, numero_documento)
        
        def configurar_estilo_cabecalho(style, size):
            style.font.color.rgb = RGBColor(0, 0, 0)
            style.font.bold = True
            style.font.size = Pt(size)
            style.paragraph_format.space_after = Pt(6)

        configurar_estilo_cabecalho(doc.styles['Heading 1'], 16)
        configurar_estilo_cabecalho(doc.styles['Heading 2'], 14)

        adicionar_secao(doc, "1. OBJETIVO", objetivo)
        adicionar_secao(doc, "2. RESPONSABILIDADES", responsabilidades)
        adicionar_secao(doc, "3. CONCEITOS E SIGLAS", conceitosSiglas)
        adicionar_secao(doc, "4. DIRETRIZES/PROCEDIMENTOS", diretrizes)
        
        doc.add_heading("5. ANEXOS", level=2)
        if request.files and 'anexos' in request.files:
            uploaded_files = request.files.getlist('anexos')
            for file_storage in uploaded_files:
                doc.add_paragraph(f"Arquivo: {file_storage.filename}")
                try:
                    para_img = doc.add_paragraph()
                    run_img = para_img.add_run()
                    with tempfile.NamedTemporaryFile(delete=False, suffix=f".{file_storage.filename.split('.')[-1]}") as tmp:
                        file_storage.save(tmp.name)
                        tmp_path = tmp.name
                    
                    run_img.add_picture(tmp_path, width=Inches(5))
                    os.unlink(tmp_path)
                except Exception as e:
                    doc.add_paragraph(f"Erro ao inserir imagem {file_storage.filename}: {str(e)}")
        else:
            doc.add_paragraph("Nenhum anexo de imagem adicionado.")

        adicionar_secao(doc, "6. DOCUMENTOS COMPLEMENTARES", documentosComplementares)
        adicionar_secao(doc, "7. REFERÊNCIAS", referencias)

        doc.add_heading("8. CONTROLE DE REVISÃO", level=2)
        if revisoes:
            table = doc.add_table(rows=1, cols=4)
            table.style = "Table Grid"
            hdr_cells = table.rows[0].cells
            for i, texto in enumerate(["Revisão", "Data", "Responsável", "Alteração"]):
                hdr_cells[i].text = texto
                hdr_cells[i].paragraphs[0].runs[0].font.bold = True

            for revisao in revisoes:
                row = table.add_row().cells
                row[0].text = numero_documento
                row[1].text = revisao.get("data", "")
                row[2].text = revisao.get("responsavel", "")
                row[3].text = revisao.get("alteracao", "")
        else:
            doc.add_paragraph("Nenhuma revisão registrada.")

        # Salva o documento localmente
        folder_path = os.path.join(BASE_PATH, pasta_nome)
        if not os.path.exists(folder_path):
            os.makedirs(folder_path)

        filename_docx = f"PSG-{sigla_pasta}-{tema_sigla}-{numero_documento}.docx"
        caminho_docx = os.path.join(folder_path, filename_docx)
        doc.save(caminho_docx)

        # Salva os dados no banco de dados
        anexos_nomes = [f.filename for f in request.files.getlist('anexos')]
        
        conn = get_db_connection()
        conn.execute("""
            INSERT INTO documentos (pasta_id, titulo, objetivo, responsaveis, conceitos_siglas, diretrizes, documentos_complementares, referencias, revisoes_json, anexos_json, data_criacao, tema_sigla)
            VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
        """, (
            pasta_id, titulo, objetivo, responsabilidades, conceitosSiglas, diretrizes,
            documentosComplementares, referencias, json.dumps(revisoes), json.dumps(anexos_nomes), datetime.now().strftime("%Y-%m-%d"), tema_sigla
        ))
        
        conn.execute("""
            INSERT INTO status (pasta_id, pasta_name, status, email)
            VALUES (?, ?, ?, ?)
        """, (
            pasta_id, pasta_nome, "Pendente", email
        ))
        conn.commit()
        conn.close()

        # Envia o documento para o SharePoint
        sharepoint_folder = "Pendentes"
        sharepoint_success = send_to_sharepoint(caminho_docx, filename_docx, sharepoint_folder)
        
        # Envia o e-mail de notificação (apenas se o upload para o SharePoint for bem-sucedido)
        if sharepoint_success:
            corpo_email = f"Olá!\n\nSeu documento PSG foi criado e está pendente de aprovação.\n\nDetalhes do Documento:\n- Título: {titulo}\n- Pasta: {pasta_nome}\n\nAguarde o retorno da gestora de LGPD. Obrigado!"
            send_notification_email(email, f"Seu PSG está pendente de aprovação.", corpo_email)

        # Envia o arquivo de volta para o front-end
        doc_io = io.BytesIO()
        doc.save(doc_io)
        doc_io.seek(0)
        
        return send_file(
            doc_io,
            mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            as_attachment=True,
            download_name=filename_docx
        )
    
    except Exception as e:
        print(f"Erro na rota gerar_documento: {e}")
        return jsonify({"error": f"Ocorreu um erro no servidor: {str(e)}"}), 500

@app.route('/documentos_por_pasta/<int:pasta_id>', methods=['GET'])
def get_documentos_por_pasta(pasta_id):
    """Retorna os dados dos documentos para uma pasta específica."""
    conn = get_db_connection()
    docs = conn.execute("""
        SELECT 
            d.id, d.titulo, d.objetivo, d.responsaveis, d.conceitos_siglas, d.diretrizes, d.documentos_complementares, d.referencias, 
            d.revisoes_json, d.anexos_json, d.data_criacao, d.tema_sigla,
            p.id AS pasta_id, p.nome AS pasta_nome, p.sigla AS pasta_sigla
        FROM documentos d
        JOIN pastas p ON d.pasta_id = p.id
        WHERE d.pasta_id = ?
    """, (pasta_id,)).fetchall()
    conn.close()

    if not docs:
        return jsonify([])

    docs_list = []
    for doc in docs:
        doc_dict = dict(doc)
        
        # Desserializa os campos JSON de forma segura
        doc_dict['revisoes'] = json.loads(doc_dict['revisoes_json'] or '[]')
        doc_dict['anexos'] = json.loads(doc_dict['anexos_json'] or '[]')
        
        # Constrói o nome do arquivo
        sigla_pasta = doc_dict['pasta_sigla']
        tema_sigla = doc_dict['tema_sigla']
        numero_documento = doc_dict['id']
        doc_dict['full_filename'] = f"PSG-{sigla_pasta}-{tema_sigla}-{numero_documento:02d}.docx"

        # Formata o objeto para o front-end
        doc_dict['folder'] = {
            'id': doc_dict['pasta_id'],
            'nome': doc_dict['pasta_nome'],
            'sigla': doc_dict['pasta_sigla']
        }
        
        del doc_dict['revisoes_json']
        del doc_dict['anexos_json']
        del doc_dict['pasta_id']
        del doc_dict['pasta_nome']
        del doc_dict['pasta_sigla']
        
        docs_list.append(doc_dict)
    
    return jsonify(docs_list)


@app.route('/listar_arquivos/<path:pasta_nome>', methods=['GET'])
def listar_arquivos(pasta_nome):
    """Retorna uma lista de nomes de arquivos para uma pasta específica."""
    
    base_dir = os.path.join(os.getcwd(), BASE_PATH)
    
    pastas_existentes = [d for d in os.listdir(base_dir) if d.lower() == pasta_nome.lower() and os.path.isdir(os.path.join(base_dir, d))]
    
    if not pastas_existentes:
        return jsonify({"error": f"Pasta '{pasta_nome}' não encontrada."}), 404
        
    pasta_correta = pastas_existentes[0]
    folder_path = os.path.join(base_dir, pasta_correta)
        
    arquivos = [f for f in os.listdir(folder_path) if os.path.isfile(os.path.join(folder_path, f))]
    
    return jsonify({"arquivos": arquivos})

@app.route('/download/<path:pasta_nome>/<path:filename>', methods=['GET'])
def download_documento(pasta_nome, filename):
    """Permite o download de um documento específico."""
    file_path = os.path.join(BASE_PATH, pasta_nome, filename)
    if os.path.exists(file_path):
        return send_file(
            file_path,
            as_attachment=True,
            download_name=os.path.basename(file_path)
        )
    return jsonify({"error": "Arquivo não encontrado."}), 404

@app.route('/atualizar_documento/<int:doc_id>', methods=['PUT'])
def atualizar_documento(doc_id):
    """
    Recebe os dados editados de um documento, atualiza o registro no banco
    e gera uma nova versão do arquivo Word.
    """
    try:
        conn = get_db_connection()
        
        # Busca o documento original pelo ID
        original_doc = conn.execute("""
            SELECT 
                d.*, p.nome AS pasta_nome, p.sigla AS sigla_pasta
            FROM documentos d
            JOIN pastas p ON d.pasta_id = p.id
            WHERE d.id = ?
        """, (doc_id,)).fetchone()
        
        if not original_doc:
            conn.close()
            return jsonify({"error": "Documento não encontrado para edição."}), 404
            
        data = request.form
        
        titulo = data.get('titulo', '')
        tema_sigla = data.get('tema', '')
        revisoes = json.loads(data.get('revisoes', '[]'))
        email = data.get('email', '')
        
        objetivo = data.get('objetivo', '')
        responsabilidades = data.get('responsabilidades', '')
        conceitosSiglas = data.get('conceitosSiglas', '')
        diretrizes = data.get('diretrizes', '')
        documentosComplementares = data.get('documentosComplementares', '')
        referencias = data.get('referencias', '')
        
        pasta_nome = original_doc['pasta_nome']
        sigla_pasta = original_doc['sigla_pasta']
        pasta_id = original_doc['pasta_id']
        
        # Recupera o nome de arquivo original
        filename_docx = f"PSG-{sigla_pasta}-{tema_sigla}-{original_doc['id']:02d}.docx"
        
        cursor = conn.cursor()
        cursor.execute("""
            UPDATE documentos SET
                titulo = ?, objetivo = ?, responsaveis = ?, conceitos_siglas = ?,
                diretrizes = ?, documentos_complementares = ?, referencias = ?,
                revisoes_json = ?, tema_sigla = ?
            WHERE id = ?
        """, (
            titulo, objetivo, responsabilidades, conceitosSiglas, diretrizes,
            documentosComplementares, referencias, json.dumps(revisoes), tema_sigla, doc_id
        ))
        
        cursor.execute("""
            INSERT INTO status (pasta_id, pasta_name, status, email)
            VALUES (?, ?, ?, ?)
        """, (
            pasta_id, pasta_nome, "Pendente", email
        ))
        conn.commit()
        conn.close()
        
        numero_documento = original_doc['id']
        
        doc = Document()
        doc = criar_cabecalho(doc, pasta_nome, sigla_pasta, tema_sigla, numero_documento)
        
        def configurar_estilo_cabecalho(style, size):
            style.font.color.rgb = RGBColor(0, 0, 0)
            style.font.bold = True
            style.font.size = Pt(size)
            style.paragraph_format.space_after = Pt(6)

        configurar_estilo_cabecalho(doc.styles['Heading 1'], 16)
        configurar_estilo_cabecalho(doc.styles['Heading 2'], 14)

        adicionar_secao(doc, "1. OBJETIVO", objetivo)
        adicionar_secao(doc, "2. RESPONSABILIDADES", responsabilidades)
        adicionar_secao(doc, "3. CONCEITOS E SIGLAS", conceitosSiglas)
        adicionar_secao(doc, "4. DIRETRIZES/PROCEDIMENTOS", diretrizes)
        
        doc.add_heading("5. ANEXOS", level=2)
        doc.add_paragraph("Nenhum anexo de imagem adicionado.")

        adicionar_secao(doc, "6. DOCUMENTOS COMPLEMENTARES", documentosComplementares)
        adicionar_secao(doc, "7. REFERÊNCIAS", referencias)

        adicionar_secao(doc, "8. CONTROLE DE REVISÃO", json.dumps(revisoes))
        
        caminho_docx = os.path.join(BASE_PATH, pasta_nome, filename_docx)
        doc.save(caminho_docx)
        
        doc_io = io.BytesIO()
        doc.save(doc_io)
        doc_io.seek(0)
        
        return send_file(
            doc_io,
            mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            as_attachment=True,
            download_name=filename_docx
        )

    except Exception as e:
        print(f"Erro na rota atualizar_documento: {e}")
        return jsonify({"error": f"Ocorreu um erro no servidor: {str(e)}"}), 500

# --- Execução do servidor ---
if __name__ == '__main__':
    if not os.path.exists(BASE_PATH):
        os.makedirs(BASE_PATH)
    app.run(host='0.0.0.0', debug=True)